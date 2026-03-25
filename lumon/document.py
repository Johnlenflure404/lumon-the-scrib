"""
Document ingestion — universal text extraction from multiple file formats.

Supported formats:
  - TXT: Direct UTF-8 text reading
  - MD:  Direct UTF-8 text reading (existing workflow)
  - PDF: Digital text extraction via PyMuPDF + OCR fallback for scanned pages
  - DOCX: Text + table extraction via python-docx
  - Images: Direct OCR via GLM-OCR
"""

import io
from typing import Callable

from PIL import Image


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text from a TXT file (UTF-8)."""
    return file_bytes.decode("utf-8", errors="replace")


def extract_text_from_md(file_bytes: bytes) -> str:
    """Extract text from a Markdown file (UTF-8). Identical to TXT."""
    return file_bytes.decode("utf-8", errors="replace")


def extract_from_pdf(
    file_bytes: bytes,
    ocr_func: Callable[[bytes, str], str] | None = None,
    progress_callback: Callable[[int, int, str], None] | None = None,
) -> str:
    """
    Extract text from a PDF file.

    Smart routing:
      - Pages with extractable text → direct extraction via PyMuPDF
      - Pages with no text (scanned/image PDFs) → render to image → OCR

    Args:
        file_bytes: Raw PDF bytes
        ocr_func: Callable(image_bytes, prompt_mode) -> str. Required for scanned pages.
        progress_callback: Optional callable(current_page, total_pages, status) for progress

    Returns:
        Combined text from all pages with page markers
    """
    import fitz  # PyMuPDF

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    total_pages = len(doc)
    pages_text = []

    for page_num in range(total_pages):
        if progress_callback:
            progress_callback(page_num, total_pages, f"Page {page_num + 1}/{total_pages}")

        page = doc[page_num]

        # Try to extract text directly
        text = page.get_text("text").strip()

        if text and len(text) > 20:
            # Page has extractable digital text
            pages_text.append(text)
        elif ocr_func:
            # Page appears to be scanned/image — render and OCR
            # Render at 300 DPI for good OCR quality
            mat = fitz.Matrix(300 / 72, 300 / 72)
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")

            ocr_text = ocr_func(img_bytes, "text")
            if ocr_text:
                pages_text.append(ocr_text)
            else:
                pages_text.append(f"[Page {page_num + 1}: OCR produced no text]")
        else:
            # No OCR available — note the empty page
            if page.get_images():
                pages_text.append(
                    f"[Page {page_num + 1}: Contains images but OCR is not configured]"
                )
            else:
                pages_text.append(f"[Page {page_num + 1}: Empty]")

    doc.close()

    if progress_callback:
        progress_callback(total_pages, total_pages, "Done")

    # Join pages with clear separators
    return "\n\n---\n\n".join(pages_text)


def extract_from_docx(
    file_bytes: bytes,
    ocr_func: Callable[[bytes, str], str] | None = None,
) -> str:
    """
    Extract text from a DOCX file.

    Extracts:
      - Paragraphs (preserving heading levels)
      - Tables (converted to Markdown tables)
      - Embedded images → OCR (if ocr_func provided)

    Args:
        file_bytes: Raw DOCX bytes
        ocr_func: Optional callable(image_bytes, prompt_mode) -> str

    Returns:
        Text in Markdown-like format
    """
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    doc = Document(io.BytesIO(file_bytes))
    parts = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            # Paragraph
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break
            if para is None:
                continue

            text = para.text.strip()
            if not text:
                parts.append("")
                continue

            # Map heading styles to Markdown
            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.split()[-1])
                    parts.append(f"{'#' * level} {text}")
                except (ValueError, IndexError):
                    parts.append(text)
            else:
                parts.append(text)

        elif tag == "tbl":
            # Table — convert to Markdown
            tbl = None
            for t in doc.tables:
                if t._element is element:
                    tbl = t
                    break
            if tbl is None:
                continue

            rows = []
            for row in tbl.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")

            if rows:
                # Add header separator after first row
                header = rows[0]
                num_cols = header.count("|") - 1
                separator = "| " + " | ".join(["---"] * num_cols) + " |"
                parts.append(header)
                parts.append(separator)
                parts.extend(rows[1:])

    # Handle embedded images if OCR is available
    if ocr_func:
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    img_bytes = rel.target_part.blob
                    ocr_text = ocr_func(img_bytes, "text")
                    if ocr_text:
                        parts.append(f"\n[Embedded Image OCR]\n{ocr_text}")
                except Exception:
                    pass

    return "\n\n".join(parts)


def extract_from_image(
    image_bytes: bytes,
    ocr_func: Callable[[bytes, str], str],
    prompt_mode: str = "text",
) -> str:
    """
    Extract text from an image via OCR.

    Args:
        image_bytes: Raw image bytes
        ocr_func: Callable(image_bytes, prompt_mode) -> str
        prompt_mode: OCR prompt mode ("text", "formula", "table")

    Returns:
        Recognized text
    """
    return ocr_func(image_bytes, prompt_mode)


def image_to_png_bytes(image_bytes: bytes) -> bytes:
    """Convert any image format to PNG bytes for consistent OCR input."""
    img = Image.open(io.BytesIO(image_bytes))
    if img.mode == "RGBA":
        # Convert RGBA to RGB (white background)
        background = Image.new("RGB", img.size, (255, 255, 255))
        background.paste(img, mask=img.split()[3])
        img = background
    elif img.mode != "RGB":
        img = img.convert("RGB")

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def ingest(
    file_bytes: bytes,
    file_type: str,
    ocr_func: Callable[[bytes, str], str] | None = None,
    progress_callback: Callable | None = None,
) -> str:
    """
    Universal document ingestion router.

    Args:
        file_bytes: Raw file bytes
        file_type: File extension (e.g., "md", "txt", "pdf", "docx", "png", "jpg")
        ocr_func: OCR callable for image/scanned content
        progress_callback: Progress callback for multi-page documents

    Returns:
        Extracted text ready for translation
    """
    file_type = file_type.lower().lstrip(".")

    if file_type in ("md", "markdown"):
        return extract_text_from_md(file_bytes)

    elif file_type in ("txt", "text"):
        return extract_text_from_txt(file_bytes)

    elif file_type == "pdf":
        return extract_from_pdf(file_bytes, ocr_func, progress_callback)

    elif file_type in ("docx",):
        return extract_from_docx(file_bytes, ocr_func)

    elif file_type in ("png", "jpg", "jpeg", "bmp", "tiff", "tif", "webp", "gif"):
        if ocr_func is None:
            raise ValueError("OCR function required for image files")
        # Normalize to PNG for consistent OCR
        png_bytes = image_to_png_bytes(file_bytes)
        return extract_from_image(png_bytes, ocr_func)

    else:
        raise ValueError(f"Unsupported file type: {file_type}")
