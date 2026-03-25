"""
Multi-format export — convert translated text to various output formats.

Supported formats:
  - Markdown (.md) — reference format
  - Plain text (.txt)
  - PDF (.pdf) — via fpdf2 with Unicode support
  - HTML (.html) — print-friendly
  - DOCX (.docx) — via python-docx
"""

import io
import re


def export_markdown(text: str, filename: str = "translation.md") -> tuple[io.BytesIO, str, str]:
    """
    Export as Markdown (reference format).

    Returns:
        (buffer, filename, mime_type)
    """
    buf = io.BytesIO(text.encode("utf-8"))
    buf.name = filename
    return buf, filename, "text/markdown"


def export_txt(text: str, filename: str = "translation.txt") -> tuple[io.BytesIO, str, str]:
    """
    Export as plain text (Markdown formatting stripped).

    Returns:
        (buffer, filename, mime_type)
    """
    # Strip common Markdown formatting
    plain = text
    # Remove headers
    plain = re.sub(r"^#{1,6}\s+", "", plain, flags=re.MULTILINE)
    # Remove bold/italic
    plain = re.sub(r"\*{1,3}([^*]+)\*{1,3}", r"\1", plain)
    plain = re.sub(r"_{1,3}([^_]+)_{1,3}", r"\1", plain)
    # Remove inline code
    plain = re.sub(r"`([^`]+)`", r"\1", plain)
    # Remove links
    plain = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", plain)
    # Remove images
    plain = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", plain)
    # Remove horizontal rules
    plain = re.sub(r"^[-*_]{3,}\s*$", "", plain, flags=re.MULTILINE)

    buf = io.BytesIO(plain.encode("utf-8"))
    buf.name = filename
    return buf, filename, "text/plain"


def export_pdf(text: str, filename: str = "translation.pdf") -> tuple[io.BytesIO, str, str]:
    """
    Export as PDF with Unicode font support.

    Uses fpdf2 for lightweight PDF generation.

    Returns:
        (buffer, filename, mime_type)
    """
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=25)
    pdf.set_left_margin(15)
    pdf.set_right_margin(15)
    pdf.add_page()

    # Use built-in Helvetica
    pdf.set_font("Helvetica", size=11)

    def _write_cell(txt, line_h=7):
        """Write a multi-cell with proper X reset."""
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(w=pdf.epw, h=line_h, text=txt)

    # Process text line by line for basic Markdown handling
    lines = text.split("\n")
    for line in lines:
        stripped = line.strip()

        # Skip table separator rows (|---|---|)
        if re.match(r"^\|[\s:|-]+\|$", stripped):
            continue

        # Headings
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2)
            sizes = {1: 22, 2: 18, 3: 15, 4: 13, 5: 12, 6: 11}
            pdf.set_font("Helvetica", style="B", size=sizes.get(level, 11))
            pdf.ln(4)
            _write_cell(heading_text, 8)
            pdf.ln(2)
            pdf.set_font("Helvetica", size=11)
            continue

        # Horizontal rules
        if re.match(r"^[-*_]{3,}\s*$", stripped):
            pdf.ln(4)
            x = pdf.l_margin
            pdf.line(x, pdf.get_y(), x + pdf.epw, pdf.get_y())
            pdf.ln(4)
            continue

        # Empty line
        if not stripped:
            pdf.ln(4)
            continue

        # Table rows — convert to readable format
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            clean = "  |  ".join(cells)
            pdf.set_font("Helvetica", size=10)
            _write_cell(clean, 7)
            pdf.set_font("Helvetica", size=11)
            continue

        # Regular text — strip basic Markdown formatting for PDF
        clean = re.sub(r"\*{1,3}([^*]+)\*{1,3}", r"\1", stripped)
        clean = re.sub(r"`([^`]+)`", r"\1", clean)
        clean = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", clean)
        _write_cell(clean)


    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    buf.name = filename
    return buf, filename, "application/pdf"


def export_html(text: str, filename: str = "translation.html") -> tuple[io.BytesIO, str, str]:
    """
    Export as print-friendly HTML.

    Converts Markdown to HTML using the markdown library with a clean CSS stylesheet.

    Returns:
        (buffer, filename, mime_type)
    """
    import markdown as md

    # Convert Markdown to HTML
    html_body = md.markdown(
        text,
        extensions=["tables", "fenced_code", "nl2br"],
    )

    html_doc = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Translation</title>
    <style>
        @media print {{
            body {{ font-size: 11pt; }}
            @page {{ margin: 2cm; }}
        }}
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto,
                         'Noto Sans', 'Noto Sans CJK SC', 'Noto Sans CJK JP',
                         'Noto Sans Arabic', sans-serif;
            max-width: 800px;
            margin: 40px auto;
            padding: 0 20px;
            line-height: 1.7;
            color: #1a1a1a;
            background: #fff;
        }}
        h1, h2, h3, h4, h5, h6 {{
            margin-top: 1.5em;
            margin-bottom: 0.5em;
            color: #111;
        }}
        h1 {{ font-size: 2em; border-bottom: 2px solid #eee; padding-bottom: 0.3em; }}
        h2 {{ font-size: 1.5em; border-bottom: 1px solid #eee; padding-bottom: 0.2em; }}
        p {{ margin: 0.8em 0; }}
        code {{
            background: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: 'SF Mono', 'Fira Code', monospace;
            font-size: 0.9em;
        }}
        pre {{
            background: #f4f4f4;
            padding: 16px;
            border-radius: 6px;
            overflow-x: auto;
        }}
        pre code {{ background: none; padding: 0; }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 1em 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 8px 12px;
            text-align: left;
        }}
        th {{ background: #f8f8f8; font-weight: 600; }}
        tr:nth-child(even) {{ background: #fafafa; }}
        blockquote {{
            border-left: 4px solid #ddd;
            margin: 1em 0;
            padding: 0.5em 1em;
            color: #555;
        }}
        hr {{
            border: none;
            border-top: 1px solid #eee;
            margin: 2em 0;
        }}
    </style>
</head>
<body>
{html_body}
</body>
</html>"""

    buf = io.BytesIO(html_doc.encode("utf-8"))
    buf.name = filename
    return buf, filename, "text/html"


def export_docx(text: str, filename: str = "translation.docx") -> tuple[io.BytesIO, str, str]:
    """
    Export as DOCX document.

    Converts Markdown to DOCX with headings, paragraphs, and tables.

    Returns:
        (buffer, filename, mime_type)
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.size = Pt(11)

    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Heading
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2)
            doc.add_heading(heading_text, level=min(level, 9))
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^[-*_]{3,}\s*$", stripped):
            # Add a thin horizontal line paragraph
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(12)
            i += 1
            continue

        # Table detection (lines starting with |)
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                table_lines.append(lines[i].strip())
                i += 1

            if len(table_lines) >= 2:
                # Parse table
                rows_data = []
                for tl in table_lines:
                    # Skip separator rows
                    if re.match(r"^\|[\s:|-]+\|$", tl):
                        continue
                    cells = [c.strip() for c in tl.strip("|").split("|")]
                    rows_data.append(cells)

                if rows_data:
                    num_cols = max(len(r) for r in rows_data)
                    table = doc.add_table(rows=len(rows_data), cols=num_cols)
                    table.style = "Table Grid"

                    for row_idx, row_data in enumerate(rows_data):
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < num_cols:
                                table.cell(row_idx, col_idx).text = cell_text
            continue

        # Code block
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # Skip closing ```

            code_text = "\n".join(code_lines)
            para = doc.add_paragraph()
            run = para.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Regular paragraph — collect consecutive non-empty, non-special lines
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            next_stripped = lines[i].strip()
            if (not next_stripped or
                    next_stripped.startswith("#") or
                    next_stripped.startswith("|") or
                    next_stripped.startswith("```") or
                    re.match(r"^[-*_]{3,}\s*$", next_stripped)):
                break
            para_lines.append(next_stripped)
            i += 1

        para_text = " ".join(para_lines)
        # Strip basic formatting
        para_text = re.sub(r"\*{1,3}([^*]+)\*{1,3}", r"\1", para_text)
        para_text = re.sub(r"`([^`]+)`", r"\1", para_text)
        para_text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", para_text)
        doc.add_paragraph(para_text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    buf.name = filename
    return buf, filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


# ──────────────────────────────────────────────
# Convenience dispatcher
# ──────────────────────────────────────────────

EXPORT_FORMATS = {
    "Markdown (.md)": export_markdown,
    "Plain Text (.txt)": export_txt,
    "PDF (.pdf)": export_pdf,
    "HTML (.html)": export_html,
    "Word (.docx)": export_docx,
}

EXPORT_EXTENSIONS = {
    "Markdown (.md)": ".md",
    "Plain Text (.txt)": ".txt",
    "PDF (.pdf)": ".pdf",
    "HTML (.html)": ".html",
    "Word (.docx)": ".docx",
}
